"""
PDF to Word Converter Module
Handles the conversion logic from PDF to DOCX format using PyMuPDF and python-docx
"""

import os
from pathlib import Path
from typing import Tuple, Optional
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


class PDFToWordConverter:
    """Main converter class for PDF to Word conversion"""
    
    def __init__(self):
        self.supported_formats = ['.pdf']
        self.output_format = '.docx'
    
    def convert(self, pdf_path: str, output_path: Optional[str] = None) -> Tuple[bool, str]:
        """
        Convert PDF to Word document
        
        Args:
            pdf_path: Path to the PDF file
            output_path: Path for the output Word file (optional)
            
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            # Validate input file
            if not os.path.exists(pdf_path):
                return False, f"PDF file not found: {pdf_path}"
            
            if not pdf_path.lower().endswith('.pdf'):
                return False, "File must be a PDF"
            
            # Generate output path if not provided
            if output_path is None:
                output_path = self._generate_output_path(pdf_path)
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Perform conversion
            self._pdf_to_docx(pdf_path, output_path)
            
            # Verify output
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                return True, f"Conversion successful!\nOutput: {output_path}\nSize: {file_size} bytes"
            else:
                return False, "Conversion completed but output file not found"
                
        except Exception as e:
            return False, f"Conversion error: {str(e)}"
    
    def _pdf_to_docx(self, pdf_path: str, docx_path: str):
        """Convert PDF to DOCX using PyMuPDF and python-docx"""
        # Open PDF
        pdf_document = fitz.open(pdf_path)
        doc = Document()
        
        try:
            # Process each page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Get text from page
                text = page.get_text()
                
                if text.strip():
                    # Add text content to document
                    doc.add_paragraph(text)
                
                # Try to extract images
                image_list = page.get_images()
                if image_list:
                    for img_index in image_list:
                        xref = img_index[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        # Skip if already an RGBA image
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("png")
                        else:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            img_data = pix.tobytes("png")
                        
                        # Add image to document
                        try:
                            doc.add_picture(BytesIO(img_data), width=Inches(5))
                        except:
                            pass  # Skip if image conversion fails
                
                # Add page break between pages (except after last page)
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
        
        finally:
            pdf_document.close()
        
        # Save document
        doc.save(docx_path)
    
    def _generate_output_path(self, pdf_path: str) -> str:
        """Generate output path based on input PDF path"""
        base_path = os.path.splitext(pdf_path)[0]
        return f"{base_path}.docx"
